from aiogram.types import CallbackQuery,FSInputFile
from aiogram.fsm.context import FSMContext
from aiogram.fsm.state import State, StatesGroup

import weasyprint
import os
import fitz
from reportlab.pdfgen import canvas
from docx import Document
import subprocess
from pdf2image import convert_from_path
from pdf2docx import Converter
from odf.opendocument import OpenDocumentText,load
from odf.text import P, Span
from odf.element import Node
from odf import teletype



from app.utils.users import *
from app.utils.files import get_file


async def admin_all_user_count(call:CallbackQuery):
    users = await get_all_users()
    
    await call.message.answer(f"{len(users)} ta user bor!")


async def docs_con(call:CallbackQuery):
    try:
        data=call.data
        lst=data.split("_")
        if len(lst)!=4:
            await call.message.answer(f"The file name should not contain the _ character.")
            if os.path.exists(output_path):
                os.remove(output_path)
            if os.path.exists(input_path):
                os.remove(input_path)
            return



        file = await get_file(lst[-1])
        dtype = lst[1]
        to_dtype = lst[2]

        input_path = f"app/files/{file.file_type}/{file.file}"
        output_path = f"app/files/{file.file_type}/output_{file.id}{lst[2]}"

        file_size=os.path.getsize(input_path)
        if file_size > 5_000_000:
            usr = await get_user_by_telegram_id(call.from_user.id)
            if usr.id_pro:
                pass
            else:
                await call.message.answer("Sorry but bot cannot convert this file.\nMaximum file size is 5MB for free plan users!")
                if os.path.exists(output_path):
                    os.remove(output_path)
                if os.path.exists(input_path):
                    os.remove(input_path)
                return
# ---------------------------------------------------------------------------------------------------------------
        if dtype == ".html":
            #___________________________________________________________

            if to_dtype == ".pdf":
                weasyprint.HTML(input_path).write_pdf(output_path)
# ---------------------------------------------------------------------------------------------------------------
        if dtype == ".docx":
            #__________________________________________________________

            if to_dtype == ".txt":
                do = Document(input_path)
                with open(output_path,"w",encoding="utf-8") as txt_file:
                    for para in do.paragraphs:
                        txt_file.write(para.text + "\n")
            #___________________________________________________________

            elif to_dtype == ".pdf":
                subprocess.run([
                    "libreoffice", "--headless",
                    "--convert-to", "pdf",
                    input_path,
                    "--outdir", f"app/files/{file.file_type}"
                ])
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                if os.path.exists(f"app/files/{file.file_type}/"+ base_name + ".pdf"):
                    os.rename(f"app/files/{file.file_type}/"+ base_name + ".pdf", f"app/files/{file.file_type}/"+ f"output_{file.id}" + ".pdf")
            #_____________________________________________________________

            elif to_dtype == ".odt":
                subprocess.run([
                    "soffice", "--headless", "--convert-to", "odt", "--outdir",
                    f"app/files/{file.file_type}", input_path
                ])
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                if os.path.exists(f"app/files/{file.file_type}/"+ base_name + ".odt"):
                    os.rename(f"app/files/{file.file_type}/"+ base_name + ".odt", f"app/files/{file.file_type}/"+ f"output_{file.id}" + ".odt")
            #________________________________________________________________

            elif to_dtype == ".rtf":
                subprocess.run([
                        "soffice", "--headless", "--convert-to", "rtf", "--outdir",
                        f"app/files/{file.file_type}", input_path
                    ])
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                if os.path.exists(f"app/files/{file.file_type}/"+ base_name + ".rtf"):
                    os.rename(f"app/files/{file.file_type}/"+ base_name + ".rtf", f"app/files/{file.file_type}/"+ f"output_{file.id}" + ".rtf")
# ---------------------------------------------------------------------------------------------------------------
        if dtype == ".txt":
            #___________________________________________________________

            if to_dtype == ".pdf":
                with open(input_path, "r", encoding="utf-8") as txt_file:
                    lines = txt_file.readlines()
                c = canvas.Canvas(output_path)
                width, height = c._pagesize
                y = height - 40

                for line in lines:
                    if y < 40:
                        c.showPage()
                        y = height - 40
                    c.drawString(40, y, line.strip())
                    y -= 15
                c.save()
            #___________________________________________________________

            elif to_dtype == ".docx":
                with open(input_path, "r", encoding="utf-8") as f:
                    content = f.read()

                doc = Document()
                doc.add_paragraph(content)
                doc.save(output_path)
            #___________________________________________________________

            elif to_dtype == ".odt":
                with open(input_path, "r", encoding="utf-8") as f:
                    lines = f.readlines()
                textdoc = OpenDocumentText()
                for line in lines:
                    p = P(text=line.strip())
                    textdoc.text.addElement(p)
                textdoc.save(output_path) 
            #________________________________________________________________

            elif to_dtype == ".rtf":
                subprocess.run([
                        "soffice", "--headless", "--convert-to", "rtf", "--outdir",
                        f"app/files/{file.file_type}", input_path
                    ])
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                if os.path.exists(f"app/files/{file.file_type}/"+ base_name + ".rtf"):
                    os.rename(f"app/files/{file.file_type}/"+ base_name + ".rtf", f"app/files/{file.file_type}/"+ f"output_{file.id}" + ".rtf")
# ---------------------------------------------------------------------------------------------------------------
        if dtype == ".pdf":
            #___________________________________________________________

            if to_dtype == ".jpg":
                if file_size > 500000:
                    usr = await get_user_by_telegram_id(call.from_user.id)
                    if usr.id_pro:
                        if file_size > 2000000:
                            return await call.message.answer("Sorry but bot cannot convert this file.\nMaximum file size is 2MB!")
                    else:
                        return await call.message.answer("Sorry but bot cannot convert this file.\nMaximum file size is 500KB for free plan users!")

                images = convert_from_path(input_path,dpi=300)
                paths = []
                for i, img in enumerate(images):
                    path = f"app/files/{file.file_type}/output_{file.id}_{i}.jpg"
                    img.save(path, "JPEG")
                    paths.append(path)

                for path in paths:
                    doc = FSInputFile(path)
                    await call.message.answer_document(document=doc)
                    if os.path.exists(path):
                        os.remove(path)
                    if os.path.exists(input_path):
                        os.remove(input_path)
                return
            #___________________________________________________________


            elif to_dtype == ".txt":
                doc = fitz.open(input_path)
                with open(output_path, "w", encoding="utf-8") as f:
                    for page in doc:
                        text = page.get_text()
                        f.write(text)
                        f.write("\n" + "-" * 80 + "\n")
                doc.close()
            #___________________________________________________________

            elif to_dtype == ".docx":
                doc = Converter(input_path)
                doc.convert(output_path,start=0,end=None)
                doc.close()         
# ---------------------------------------------------------------------------------------------------------------
        if dtype == ".odt":
            #___________________________________________________________

            if to_dtype == ".pdf":
                subprocess.run([
                    "libreoffice",
                    "--headless",
                    "--convert-to", "pdf",
                    input_path,
                    "--outdir", f"app/files/{file.file_type}"
                ], stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                if os.path.exists(f"app/files/{file.file_type}/"+ base_name + ".pdf"):
                    os.rename(f"app/files/{file.file_type}/"+ base_name + ".pdf", f"app/files/{file.file_type}/"+ f"output_{file.id}" + ".pdf")
            #_____________________________________________________________

            elif to_dtype == ".txt":
                textdoc = load(input_path)
                paras = textdoc.getElementsByType(P)
                lines = [teletype.extractText(p) for p in paras]

                with open(output_path,"w",encoding="utf-8") as txt_f:
                    txt_f.write("\n".join(lines))
            #______________________________________________________________

            elif to_dtype == ".docx":
                subprocess.run([
                        "soffice", "--headless", "--convert-to", "docx", "--outdir",
                        f"app/files/{file.file_type}", input_path
                    ])
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                if os.path.exists(f"app/files/{file.file_type}/"+ base_name + ".docx"):
                    os.rename(f"app/files/{file.file_type}/"+ base_name + ".docx", f"app/files/{file.file_type}/"+ f"output_{file.id}" + ".docx")
            #________________________________________________________________

            elif to_dtype == ".rtf":
                subprocess.run([
                        "soffice", "--headless", "--convert-to", "rtf", "--outdir",
                        f"app/files/{file.file_type}", input_path
                    ])
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                if os.path.exists(f"app/files/{file.file_type}/"+ base_name + ".rtf"):
                    os.rename(f"app/files/{file.file_type}/"+ base_name + ".rtf", f"app/files/{file.file_type}/"+ f"output_{file.id}" + ".rtf")
# ---------------------------------------------------------------------------------------------------------------
        if dtype == ".rtf":
            #_______________________________________________________________

            if to_dtype in [".txt",".odt",".docx"]:
                subprocess.run([
                        "soffice", "--headless", "--convert-to", f"{to_dtype[1:]}", "--outdir",
                        f"app/files/{file.file_type}", input_path
                    ])
                base_name = os.path.splitext(os.path.basename(input_path))[0]
                if os.path.exists(f"app/files/{file.file_type}/"+ base_name + f"{to_dtype}"):
                    os.rename(f"app/files/{file.file_type}/"+ base_name + f"{to_dtype}", f"app/files/{file.file_type}/"+ f"output_{file.id}" + f"{to_dtype}")




        doc=FSInputFile(output_path)
        await call.message.answer_document(document=doc, caption=f"Your file converted from {lst[1]} to {lst[2]}")
        if os.path.exists(output_path):
            os.remove(output_path)
        if os.path.exists(input_path):
            os.remove(input_path)


    except Exception as e:
        if "No such file or directory:" in str(e):
            await call.message.answer(f"If you have selected a file before, please resend the file and convert it. Because we will automatically delete the file after converting it. For data security")
            if os.path.exists(output_path):
                os.remove(output_path)
            if os.path.exists(input_path):
                os.remove(input_path)
        elif "HTTP Client says - Request Entity Too Large" in str(e):
            await call.message.answer(f"Your file size has exceeded the telegram bots limit!")
            if os.path.exists(output_path):
                os.remove(output_path)
            if os.path.exists(input_path):
                os.remove(input_path)
        else:
            await call.message.answer(f"Sorry but bot cannot convert this file!")
            if os.path.exists(output_path):
                os.remove(output_path)
            if os.path.exists(input_path):
                os.remove(input_path)

#...






